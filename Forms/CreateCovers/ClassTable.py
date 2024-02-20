import math
import os
from Base import *
from docx import Document
from docx import oxml
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor
from Config import *
from Classes.ClassUtility import *
from docx2pdf import convert
from Classes.ClassWord import Word
from Data.WordElement import *
from Data.TableElement import *
import numpy as np


class Table:
    MY_TABLE_STYLE_FINAL = 'MyTableStyle'
    MY_TABLE_TITLE_STYLE_FINAL = 'MyPerStyleTableTitles'
    MY_PER_STYLE_FINAL = 'MyPerStyle'
    MY_CODE_HEADER_STYLE_FINAL = 'MyCode'
    MY_RELEASE_DATE_HEADER_STYLE_FINAL = 'MyReleaseDate'
    MY_SCORE_HEADER_STYLE_FINAL = 'MyScore'
    MY_FINISHED_HEADER_STYLE_FINAL = 'MyFinished'
    MY_PER_TITLE_HEADER_STYLE_FINAL = 'MyTitlePer'
    MY_ENG_TITLE_HEADER_STYLE_FINAL = 'MyTitleEng'
    MY_DUBBED_STYLE_FINAL = 'MyDubbedStyle'
    MY_SUBBED_STYLE_FINAL = 'MySubbedStyle'
    My_EMPTY_STYLE = "MyEmptyStyle"

    absolute_path = rename_file_paths_by_os(os.path.dirname(__file__))
    extras_path = "extras"
    full_path = rename_file_paths_by_os(absolute_path + extras_path)

    template = full_path + "template.docx"

    final_word_location = full_path + "ex.docx"

    final_pdf_location = full_path + "james"

    document = Document(template)

    sections = document.sections

    code_header_style = document.styles[MY_CODE_HEADER_STYLE_FINAL]
    release_date_style = document.styles[MY_RELEASE_DATE_HEADER_STYLE_FINAL]
    score_style = document.styles[MY_SCORE_HEADER_STYLE_FINAL]
    finished_style = document.styles[MY_FINISHED_HEADER_STYLE_FINAL]
    eng_title_header_style = document.styles[MY_ENG_TITLE_HEADER_STYLE_FINAL]
    per_title_header_style = document.styles[MY_PER_TITLE_HEADER_STYLE_FINAL]
    dubbed_style = document.styles[MY_DUBBED_STYLE_FINAL]
    subbed_style = document.styles[MY_SUBBED_STYLE_FINAL]
    table_style = document.styles[MY_TABLE_STYLE_FINAL]
    per_style_table_titles = document.styles[MY_TABLE_TITLE_STYLE_FINAL]
    per_style = document.styles[MY_PER_STYLE_FINAL]
    empty_style = document.styles[My_EMPTY_STYLE]

    MODE_1_PIC_WIDTH = 8.00
    MODE_1_PIC_HEIGHT = 6.55

    MODE_2_PIC_WIDTH = 4.5
    MODE_2_PIC_HEIGHT_TYPE_0 = 5.65
    MODE_2_PIC_HEIGHT_TYPE_1 = 5.65
    MODE_2_PIC_HEIGHT_TYPE_2 = 5.65
    MODE_2_PIC_HEIGHT_TYPE_3 = 5.65

    MODE_4_PIC_WIDTH = 3.97
    MODE_4_PIC_HEIGHT_TYPE_0 = 4.20
    MODE_4_PIC_HEIGHT_TYPE_1 = 4.70
    MODE_4_PIC_HEIGHT_TYPE_2 = 4.50
    MODE_4_PIC_HEIGHT_TYPE_3 = 3.75

    MODE_8_PIC_WIDTH = 1.90
    MODE_8_PIC_HEIGHT_TYPE_0 = 4.40
    MODE_8_PIC_HEIGHT_TYPE_1 = 4.95
    MODE_8_PIC_HEIGHT_TYPE_2 = 4.70
    MODE_8_PIC_HEIGHT_TYPE_3 = 4.15

    MODE_9_PIC_WIDTH = 2.60
    MODE_9_PIC_HEIGHT_TYPE_0 = 2.80
    MODE_9_PIC_HEIGHT_TYPE_1 = 3.17
    MODE_9_PIC_HEIGHT_TYPE_2 = 3.0
    MODE_9_PIC_HEIGHT_TYPE_3 = 2.40

    MODE_12_PIC_WIDTH = 1.95
    MODE_12_PIC_HEIGHT_TYPE_0 = 2.80
    MODE_12_PIC_HEIGHT_TYPE_1 = 3.17
    MODE_12_PIC_HEIGHT_TYPE_2 = 3.0
    MODE_12_PIC_HEIGHT_TYPE_3 = 2.60

    MODE_15_PIC_WIDTH = 1.52
    MODE_15_PIC_HEIGHT_TYPE_0 = 1.50
    MODE_15_PIC_HEIGHT_TYPE_1 = 1.85
    MODE_15_PIC_HEIGHT_TYPE_2 = 1.72
    MODE_15_PIC_HEIGHT_TYPE_3 = 2.80

    MODE_16_PIC_WIDTH = 2.00
    MODE_16_PIC_HEIGHT_TYPE_0 = 2.10
    MODE_16_PIC_HEIGHT_TYPE_1 = 2.10
    MODE_16_PIC_HEIGHT_TYPE_2 = 2.10
    MODE_16_PIC_HEIGHT_TYPE_3 = 1.80

    MODE_18_PIC_WIDTH = 1.25
    MODE_18_PIC_HEIGHT_TYPE_0 = 2.70
    MODE_18_PIC_HEIGHT_TYPE_1 = 3.20
    MODE_18_PIC_HEIGHT_TYPE_2 = 3.00
    MODE_18_PIC_HEIGHT_TYPE_3 = 2.40

    MODE_20_PIC_WIDTH = 1.50
    MODE_20_PIC_HEIGHT_TYPE_0 = 2.00
    MODE_20_PIC_HEIGHT_TYPE_1 = 2.35
    MODE_20_PIC_HEIGHT_TYPE_2 = 2.20
    MODE_20_PIC_HEIGHT_TYPE_3 = 1.75

    MODE_2_dub_sub_cell_width = 0.11
    MODE_4_dub_sub_cell_width = 0.11
    MODE_9_dub_sub_cell_width = 0.16
    MODE_12_dub_sub_cell_width = 0.13
    MODE_16_dub_sub_cell_width = 0.11
    MODE_20_dub_sub_cell_width = 0.11

    film_type = 0

    DUBBED_MODE_FINAL = 1
    SUBBED_MODE_FINAL = 1

    HAS_RELEASE_IN_COVER = 1
    HAS_SCORE_IN_COVER = 1
    HAS_TV_SERIES_FINISHED_IN_COVER = 1

    EXPORT_DOCX = 1
    EXPORT_PDF = 0

    def __init__(self, films_array, mode, film_type, has_release_in_cover,
                 has_score_in_cover, tv_series_finished, word_or_pdf_name=None):
        self.document._body.clear_content()

        self.final_word_location = self.full_path + word_or_pdf_name + ".docx"
        self.final_pdf_location = self.full_path + word_or_pdf_name + ".pdf"

        self.film_type = film_type

        # self.DUBBED_MODE_FINAL = checkbox_dubbed
        # self.SUBBED_MODE_FINAL = checkbox_subbed
        self.HAS_RELEASE_IN_COVER = has_release_in_cover
        self.HAS_SCORE_IN_COVER = has_score_in_cover
        self.HAS_TV_SERIES_FINISHED_IN_COVER = tv_series_finished

        self.initial_style_colors()

        if mode == 1:
            self.initial_style_sizes(code=18, release_date=18, score=15, finished=18, eng_title=16, per_title=16,
                                     dubbed=14, table_titles=14, table_contents=13)
            self.poster_one_part(films_array)
        elif mode == 2:
            self.initial_style_sizes(code=14, release_date=12, score=12, finished=12, eng_title=12, per_title=12,
                                     dubbed=9, table_titles=12, table_contents=10)
            self.poster_two_part(films_array)
        # elif mode == 2 and film_type == 2:
        #     self.initial_style_sizes(code=15, release_date=15,score=15, finished=15, eng_title=15, per_title=15,
        #                              dubbed=12, table_titles=15, table_contents=13)
        #     self.poster_two_part(films_array)
        elif mode == 4:
            self.initial_style_sizes(code=15, release_date=12, score=12, finished=13, eng_title=12, per_title=13,
                                     dubbed=12, table_titles=13, table_contents=13)
            self.poster_four_part(films_array)

        elif mode == 8:
            self.initial_style_sizes(code=8, release_date=8, score=8, finished=8, eng_title=9, per_title=9,
                                     dubbed=7, table_titles=8, table_contents=8)
            self.poster_eight_part(films_array)

        elif mode == 9:
            self.initial_style_sizes(code=9, release_date=8, score=8, finished=8, eng_title=9, per_title=9,
                                     dubbed=7.0, table_titles=8, table_contents=8)
            self.poster_nine_part(films_array)

        elif mode == 12:
            self.initial_style_sizes(code=8, release_date=8, score=8, finished=8, eng_title=8, per_title=8,
                                     dubbed=6.0, table_titles=8, table_contents=8)
            self.poster_twelve_part(films_array)

        elif mode == 15:
            self.initial_style_sizes(code=6, release_date=6, score=6, finished=6, eng_title=7, per_title=7,
                                     dubbed=6.0, table_titles=6, table_contents=6)
            self.poster_fifteen_part(films_array)

        elif mode == 16:
            self.initial_style_sizes(code=7, release_date=6, score=6, finished=7, eng_title=7, per_title=7,
                                     dubbed=5.0, table_titles=5, table_contents=5)
            self.poster_sixteen_part(films_array)

        elif mode == 18:
            self.initial_style_sizes(code=7, release_date=6, score=6, finished=5, eng_title=7, per_title=7,
                                     dubbed=5.0, table_titles=5, table_contents=5)
            self.poster_eighteen_part(films_array)

        elif mode == 20:
            self.initial_style_sizes(code=6, release_date=6, score=6, finished=5, eng_title=7, per_title=7,
                                     dubbed=5, table_titles=5, table_contents=5)
            self.poster_twenty_part(films_array)

    def initial_style_sizes(self, code, release_date, score, finished, eng_title, per_title, dubbed, table_titles,
                            table_contents):
        self.code_header_style.font.size = Pt(code)
        self.release_date_style.font.size = Pt(release_date)
        self.score_style.font.size = Pt(score)
        self.finished_style.font.size = Pt(finished)
        self.eng_title_header_style.font.size = Pt(eng_title)
        self.per_title_header_style.font.size = Pt(per_title)
        self.dubbed_style.font.size = Pt(dubbed)
        self.subbed_style.font.size = Pt(dubbed)
        self.per_style_table_titles.font.size = Pt(table_titles)
        self.per_style.font.size = Pt(table_contents)
        self.empty_style.font.size = Pt(0)

    def initial_style_colors(self):
        self.code_header_style.font.color.rgb = CustomColorRGB.BLUE
        self.release_date_style.font.color.rgb = CustomColorRGB.RED
        self.score_style.font.color.rgb = CustomColorRGB.RED
        self.finished_style.font.color.rgb = CustomColorRGB.DarkGreen
        self.eng_title_header_style.font.color.rgb = CustomColorRGB.RED
        self.per_title_header_style.font.color.rgb = CustomColorRGB.RED
        self.dubbed_style.font.color.rgb = CustomColorRGB.ORANGE
        self.subbed_style.font.color.rgb = CustomColorRGB.ORANGE
        self.per_style_table_titles.font.color.rgb = CustomColorRGB.RED

    def set_margins(self):
        for section in self.sections:
            section.top_margin = Inches(0.1)
            section.bottom_margin = Inches(0.1)
            section.left_margin = Inches(0.1)
            section.right_margin = Inches(0.1)

            # section.left_margin = Inches(0.15)
            # section.right_margin = Inches(0.15)

    @staticmethod
    def is_dubbed_or_subbed(film_path):
        per_sub_state = has_per_subtitle(film_path)

        is_subbed = 0
        is_dubbed = 0

        if per_sub_state == 1 or per_sub_state == 3:
            is_subbed = 1

        if per_sub_state == 2 or per_sub_state == 3:
            is_dubbed = 1

    @staticmethod
    def get_data_based_on_type_and_part(film_type, has_score_in_cover, current_film, film_info, poster_code,
                                        per_title_element, eng_title_element, has_dub_sub_title):
        has_eng_title = False
        has_per_title = False

        if film_type == 0 or film_type == 3:
            if has_score_in_cover:
                dict = {
                    ": بازیگران": current_film.actors,
                    ": ژانر": current_film.per_genre,
                    ": امتیاز": current_film.score,
                    ": خلاصه داستان": current_film.per_info
                }
            else:
                dict = {
                    ": بازیگران": current_film.actors,
                    ": ژانر": current_film.per_genre,
                    ": خلاصه داستان": current_film.per_info
                }

            film_info.append(dict)
            has_per_title = True
            has_eng_title = True
            has_dub_sub_title.append(1)

        elif film_type == 1:
            dict = {
                ": کارگردان": current_film.director,
                ": بازیگران": current_film.actors,
                ": ژانر": current_film.per_genre,
                ": خلاصه داستان": current_film.per_info
            }

            film_info.append(dict)
            has_per_title = True

        elif film_type == 2:
            dict = {
                ": خلاصه داستان": current_film.per_info
            }

            film_info.append(dict)
            has_per_title = True
            has_eng_title = True

        if has_eng_title:
            if current_film.eng_title:
                word_element = WordElement(current_film.eng_title, Table.eng_title_header_style,
                                           alignment=1)

                eng_title_element.append(word_element)

        if has_per_title:
            if current_film.per_title:
                word_element = WordElement(current_film.per_title, Table.per_title_header_style,
                                           alignment=1)

                per_title_element.append(word_element)

    def poster_one_part(self, films_array):

        self.set_margins()

        for i in films_array:
            Print.print_green(get_last_file_or_folder_from_path(i[0]))
            pic_location = rename_file_paths_for_files(i[0] + "/" + "Folder.jpg")

            film_path = i[0]

            current_film = i[1]

            Word.create_paragraph_image(parent=self.document, pic_dir=pic_location, pic_width=self.MODE_1_PIC_WIDTH,
                                        pic_height=self.MODE_1_PIC_HEIGHT)

            poster_code = "(" + (File.trim_poster_code(get_last_file_or_folder_from_path(i[0]))) + ")"

            film_info_dict = []
            per_title_element = []
            eng_title_element = []
            has_dub_sub_title = []

            self.get_data_based_on_type_and_part(self.film_type, self.HAS_SCORE_IN_COVER,
                                                 current_film, film_info_dict, poster_code,
                                                 per_title_element, eng_title_element, has_dub_sub_title)

            # ----------------------- Poster Code & Release Date in Table ----------------------- #
            if has_dub_sub_title:
                table = TableElement(self.table_style, 1)

                temp_header_array = []
                final_array = []

                is_subbed, is_dubbed = has_per_subtitle(film_path)

                temp_header_array.append(WordElement(poster_code, self.code_header_style))
                temp_header_array.append(WordElement("", self.code_header_style))

                temp_header_array.append(
                    WordElement("زیرنویس", self.subbed_style, is_checkbox=1, check_state=is_subbed))
                temp_header_array.append(
                    WordElement("دوبله", self.dubbed_style, is_checkbox=1, check_state=is_dubbed))

                temp_header_array.append(WordElement("", self.code_header_style))

                if self.HAS_RELEASE_IN_COVER:
                    temp_header_array.append(WordElement(current_film.release, self.release_date_style))
                else:
                    temp_header_array.append(WordElement("", self.release_date_style))

                final_array.append(temp_header_array)

                Word.create_simple_table(document=self.document, parent=self.document,
                                         table=table,
                                         content=final_array,
                                         empty_style=self.empty_style
                                         )

            else:
                # ----------------------- poster code alone ----------------------- #
                if poster_code:
                    Word.create_paragraph_text(parent=self.document,
                                               element=WordElement(poster_code, self.code_header_style, alignment=3)
                                               )
                # --------------------------------------------------------------------------------- #

            # Word.create_empty_paragraph(self.document, self.per_style)

            if self.film_type == 3:
                Word.add_input_checkbox_content_control(parent=self.document,
                                                        element=WordElement(
                                                            "تمام شده", self.finished_style, is_checkbox=1,
                                                            check_state=self.HAS_TV_SERIES_FINISHED_IN_COVER
                                                        ))

            if eng_title_element:
                Word.create_paragraph_text(self.document,
                                           element=eng_title_element[0])

            if per_title_element:
                Word.create_paragraph_text(self.document, element=per_title_element[0])

            score_cell_styles = [
                self.eng_title_header_style,
                self.per_style_table_titles,
                self.eng_title_header_style
            ]

            if self.film_type == 2:
                Word.create_empty_paragraph(self.document, self.per_style)
                Word.create_empty_paragraph(self.document, self.per_style)

            Word.create_table(self.document, rows=len(film_info_dict[0]), cols=2,
                              table_styles=self.table_style,
                              table_cel1_style=self.per_style_table_titles,
                              table_cel2_style=self.per_style,
                              empty_style=self.empty_style,
                              score_cell_styles=score_cell_styles,
                              content=film_info_dict[0],
                              cell1_width=1.3, cell1_height=0.60,
                              cell2_width=6.7)

        self.document.save(self.final_word_location)

        if self.EXPORT_DOCX:
            os.system(self.final_word_location)

        if self.EXPORT_PDF:
            Word.convert_to_pdf(self.final_word_location, self.final_pdf_location)

    def poster_two_part(self, films_array):
        self.set_margins()

        for i in films_array:
            Print.print_green(get_last_file_or_folder_from_path(i[0]))
            pic_location = rename_file_paths_for_files(i[0] + "/" + "Folder.jpg")

            film_path = i[0]

            current_film = i[1]

            poster_code = "(" + File.trim_poster_code(get_last_file_or_folder_from_path(i[0])) + ")"

            film_info_dict = []
            per_title_element = []
            eng_title_element = []
            has_dub_sub_title = []

            table1 = self.document.add_table(rows=1, cols=2)
            table1.style = self.table_style
            table1.alignment = 1

            cell1 = table1.rows[0].cells[1]
            # cell1.paragraphs[0].style = self.per_style_table_titles
            cell1.paragraphs[0].style = self.empty_style
            # Word.delete_paragraph(cell1.paragraphs[0])
            cell1.width = Inches(3.67)

            cell2 = table1.rows[0].cells[0]
            # par = cell2.paragraphs[0]
            # run = par.add_run()

            Word.delete_paragraph(cell2.paragraphs[0])

            if self.film_type == 0:
                Word.create_paragraph_image(parent=cell2, style=self.eng_title_header_style, pic_dir=pic_location,
                                            pic_width=self.MODE_2_PIC_WIDTH,
                                            pic_height=self.MODE_2_PIC_HEIGHT_TYPE_0,
                                            alignment=1)

            elif self.film_type == 1:
                Word.create_paragraph_image(parent=cell2, style=self.eng_title_header_style, pic_dir=pic_location,
                                            pic_width=self.MODE_2_PIC_WIDTH,
                                            pic_height=self.MODE_2_PIC_HEIGHT_TYPE_1,
                                            alignment=1)

            elif self.film_type == 2:
                Word.create_paragraph_image(parent=cell2, style=self.eng_title_header_style, pic_dir=pic_location,
                                            pic_width=self.MODE_2_PIC_WIDTH,
                                            pic_height=self.MODE_2_PIC_HEIGHT_TYPE_2,
                                            alignment=1)

            elif self.film_type == 3:
                Word.create_paragraph_image(parent=cell2, style=self.eng_title_header_style, pic_dir=pic_location,
                                            pic_width=self.MODE_2_PIC_WIDTH,
                                            pic_height=self.MODE_2_PIC_HEIGHT_TYPE_3,
                                            alignment=1)

            self.get_data_based_on_type_and_part(self.film_type, self.HAS_SCORE_IN_COVER, current_film,
                                                 film_info_dict, poster_code, per_title_element,
                                                 eng_title_element, has_dub_sub_title)

            # ----------------------- Poster Code & Release Date in Table ----------------------- #
            if has_dub_sub_title:
                table = TableElement(self.table_style, 1)

                temp_header_array = []
                final_array = []

                temp_header_array = []
                final_array = []

                is_subbed, is_dubbed = has_per_subtitle(film_path)

                temp_header_array.append(WordElement(poster_code, self.code_header_style))
                # temp_header_array.append(WordElement("", self.code_header_style))

                score = current_film.score + "/10"

                temp_header_array.append(
                    WordElement(score, self.score_style))

                temp_header_array.append(
                    WordElement("زیرنویس", self.subbed_style, is_checkbox=1, check_state=is_subbed,
                                width=self.MODE_2_dub_sub_cell_width))
                temp_header_array.append(
                    WordElement("دوبله", self.dubbed_style, is_checkbox=1, check_state=is_dubbed,
                                width=self.MODE_2_dub_sub_cell_width))

                # temp_header_array.append(WordElement("", self.code_header_style))

                if self.HAS_RELEASE_IN_COVER:
                    temp_header_array.append(WordElement(current_film.release, self.release_date_style))
                else:
                    temp_header_array.append(WordElement("", self.release_date_style))

                final_array.append(temp_header_array)

                Word.create_simple_table(document=self.document, parent=cell1,
                                         table=table,
                                         content=final_array,
                                         empty_style=self.empty_style,
                                         mode=2,
                                         film_type=self.film_type
                                         )
            else:
                Word.delete_paragraph(cell1.paragraphs[0])
                if poster_code:
                    Word.create_paragraph_text(parent=cell1,
                                               element=WordElement(poster_code, self.code_header_style, alignment=3)
                                               )
            # --------------------------------------------------------------------------------- #

            # Word.create_empty_paragraph(cell1, self.empty_style)

            if self.film_type == 3:
                Word.add_input_checkbox_content_control(parent=cell1,
                                                        element=WordElement(
                                                            "تمام شده", self.finished_style, is_checkbox=1,
                                                            check_state=self.HAS_TV_SERIES_FINISHED_IN_COVER
                                                        ))

            # Word.create_paragraph_text(parent=cell1, element=WordElement("", self.code_header_style, alignment=3))

            if eng_title_element:
                Word.create_paragraph_text(parent=cell1, element=eng_title_element[0])

            if per_title_element:
                Word.create_paragraph_text(parent=cell1, element=per_title_element[0])

            score_cell_styles = [
                self.eng_title_header_style,
                self.per_style_table_titles,
                self.eng_title_header_style
            ]

            Word.create_table(cell1, rows=len(film_info_dict[0]), cols=2,
                              table_styles=self.table_style,
                              table_cel1_style=self.per_style_table_titles,
                              table_cel2_style=self.per_style,
                              empty_style=self.empty_style,
                              score_cell_styles=score_cell_styles,
                              content=film_info_dict[0],
                              cell1_width=1.25, cell1_height=0.60,
                              cell2_width=2.27)

        self.document.save(self.final_word_location)

        if self.EXPORT_DOCX:
            os.system(self.final_word_location)

        if self.EXPORT_PDF:
            Word.convert_to_pdf(self.final_word_location, self.final_pdf_location)

    def poster_four_part(self, films_array):
        self.set_margins()

        if self.film_type == 1 or self.film_type == 2:
            self.MODE_4_PIC_HEIGHT = 4.80

        films_length = len(films_array)

        table1 = self.document.add_table(rows=math.ceil(films_length / 2), cols=2)

        counter = 0

        for i in table1.rows:
            for j in i.cells:
                if counter < films_length:
                    film = films_array[counter]

                    film_path = film[0]

                    Print.print_green(get_last_file_or_folder_from_path(film[0]))

                    pic_location = rename_file_paths_for_files(film[0] + "\\" + "Folder.jpg")
                    film_info = film[1]
                    poster_code = "(" + File.trim_poster_code(get_last_file_or_folder_from_path(film[0])) + ")"

                    # j.paragraphs[0].style = self.empty_style
                    Word.delete_paragraph(j.paragraphs[0])

                    if self.film_type == 0:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_4_PIC_WIDTH,
                                                    pic_height=self.MODE_4_PIC_HEIGHT_TYPE_0,
                                                    alignment=1)

                    elif self.film_type == 1:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_4_PIC_WIDTH,
                                                    pic_height=self.MODE_4_PIC_HEIGHT_TYPE_1,
                                                    alignment=1)

                    elif self.film_type == 2:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_4_PIC_WIDTH,
                                                    pic_height=self.MODE_4_PIC_HEIGHT_TYPE_2,
                                                    alignment=1)

                    elif self.film_type == 3:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_4_PIC_WIDTH,
                                                    pic_height=self.MODE_4_PIC_HEIGHT_TYPE_3,
                                                    alignment=1)

                    # ----------------------- Poster Code & Release Date in Table ----------------------- #
                    if self.film_type == 0 or self.film_type == 3:
                        table = TableElement(self.table_style, 1)

                        temp_header_array = []
                        final_array = []

                        is_subbed, is_dubbed = has_per_subtitle(film_path)

                        temp_header_array.append(WordElement(poster_code, self.code_header_style))
                        # temp_header_array.append(WordElement("", self.code_header_style))

                        score = film_info.score + "/10"

                        temp_header_array.append(
                            WordElement(score, self.score_style))

                        temp_header_array.append(
                            WordElement("زیرنویس", self.subbed_style, is_checkbox=1,
                                        check_state=is_subbed, width=self.MODE_4_dub_sub_cell_width))
                        temp_header_array.append(
                            WordElement("دوبله", self.dubbed_style, is_checkbox=1,
                                        check_state=is_dubbed, width=self.MODE_4_dub_sub_cell_width))

                        # temp_header_array.append(WordElement("", self.code_header_style))

                        if self.HAS_RELEASE_IN_COVER:
                            temp_header_array.append(WordElement(film_info.release, self.release_date_style))
                        else:
                            temp_header_array.append(WordElement("", self.release_date_style))

                        final_array.append(temp_header_array)

                        Word.create_simple_table(document=self.document, parent=j,
                                                 table=table,
                                                 content=final_array,
                                                 empty_style=self.empty_style,
                                                 mode=4
                                                 )

                        if self.film_type == 3:
                            Word.add_input_checkbox_content_control(parent=j,
                                                                    element=WordElement(
                                                                        "تمام شده", self.finished_style, is_checkbox=1,
                                                                        check_state=self.HAS_TV_SERIES_FINISHED_IN_COVER
                                                                    ))

                    if self.film_type == 1 or self.film_type == 2:
                        table = TableElement(self.table_style, 1)

                        temp_header_array = []
                        final_array = []

                        temp_header_array.append(WordElement(poster_code, self.code_header_style))
                        temp_header_array.append(WordElement("", self.code_header_style))
                        temp_header_array.append(WordElement("", self.code_header_style))

                        if self.HAS_RELEASE_IN_COVER:
                            temp_header_array.append(WordElement(film_info.release, self.release_date_style))
                        else:
                            temp_header_array.append(WordElement("", self.release_date_style))

                        final_array.append(temp_header_array)

                        Word.create_simple_table(document=self.document, parent=j,
                                                 table=table,
                                                 content=final_array,
                                                 empty_style=self.empty_style,
                                                 mode=4
                                                 )

                    if self.film_type == 0 or self.film_type == 2 or self.film_type == 3:
                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(film_info.eng_title, self.eng_title_header_style,
                                                                       alignment=1))

                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(film_info.per_title, self.per_title_header_style,
                                                                       alignment=1))
                    elif self.film_type == 1:
                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(film_info.per_title, self.per_title_header_style,
                                                                       alignment=1))

                    if self.film_type == 0 or self.film_type == 1 or self.film_type == 3:
                        per_genre = " , ".join(film_info.per_genre)
                        genre_text = "(" + per_genre + ")"
                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(genre_text, self.per_title_header_style,
                                                                       alignment=1))

                counter += 1

        # Word.create_paragraph_text(parent=self.document, element=WordElement(content="", style=self.empty_style))

        self.document.save(self.final_word_location)

        if self.EXPORT_DOCX:
            os.system(self.final_word_location)

        if self.EXPORT_PDF:
            Word.convert_to_pdf(self.final_word_location, self.final_pdf_location)

    def poster_eight_part(self, films_array):
        self.set_margins()

        # if self.film_type == 1 or self.film_type == 2:
        #     self.MODE_8_PIC_HEIGHT = 4.80

        films_length = len(films_array)

        table1 = self.document.add_table(rows=math.ceil(films_length / 2), cols=4)

        counter = 0

        for i in table1.rows:
            for j in i.cells:
                if counter < films_length:
                    film = films_array[counter]

                    film_path = film[0]

                    Print.print_green(get_last_file_or_folder_from_path(film[0]))

                    pic_location = rename_file_paths_for_files(film[0] + "\\" + "Folder.jpg")
                    film_info = film[1]
                    poster_code = "(" + File.trim_poster_code(get_last_file_or_folder_from_path(film[0])) + ")"

                    # j.paragraphs[0].style = self.empty_style
                    Word.delete_paragraph(j.paragraphs[0])

                    if self.film_type == 0:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_8_PIC_WIDTH,
                                                    pic_height=self.MODE_8_PIC_HEIGHT_TYPE_0,
                                                    alignment=1)

                    elif self.film_type == 1:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_8_PIC_WIDTH,
                                                    pic_height=self.MODE_8_PIC_HEIGHT_TYPE_1,
                                                    alignment=1)

                    elif self.film_type == 2:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_8_PIC_WIDTH,
                                                    pic_height=self.MODE_8_PIC_HEIGHT_TYPE_2,
                                                    alignment=1)

                    elif self.film_type == 3:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_8_PIC_WIDTH,
                                                    pic_height=self.MODE_8_PIC_HEIGHT_TYPE_3,
                                                    alignment=1)

                    # ----------------------- Poster Code & Release Date in Table ----------------------- #
                    if self.film_type == 0 or self.film_type == 3:
                        table = TableElement(self.table_style, 1)

                        temp_header_array = []
                        final_array = []

                        is_subbed, is_dubbed = has_per_subtitle(film_path)

                        temp_header_array.append(WordElement(poster_code, self.code_header_style))
                        # temp_header_array.append(WordElement("", self.code_header_style))

                        score = film_info.score + "/10"

                        temp_header_array.append(
                            WordElement(score, self.score_style))

                        temp_header_array.append(
                            WordElement("زیرنویس", self.subbed_style, is_checkbox=1,
                                        check_state=is_subbed, width=self.MODE_9_dub_sub_cell_width))
                        temp_header_array.append(
                            WordElement("دوبله", self.dubbed_style, is_checkbox=1,
                                        check_state=is_dubbed, width=self.MODE_9_dub_sub_cell_width))

                        # temp_header_array.append(WordElement("", self.code_header_style))

                        if self.HAS_RELEASE_IN_COVER:
                            temp_header_array.append(WordElement(film_info.release, self.release_date_style))
                        else:
                            temp_header_array.append(WordElement("", self.release_date_style))

                        final_array.append(temp_header_array)

                        Word.create_simple_table(document=self.document, parent=j,
                                                 table=table,
                                                 content=final_array,
                                                 empty_style=self.empty_style,
                                                 mode=4
                                                 )

                        if self.film_type == 3:
                            Word.add_input_checkbox_content_control(parent=j,
                                                                    element=WordElement(
                                                                        "تمام شده", self.finished_style, is_checkbox=1,
                                                                        check_state=self.HAS_TV_SERIES_FINISHED_IN_COVER
                                                                    ))

                    if self.film_type == 1 or self.film_type == 2:
                        table = TableElement(self.table_style, 1)

                        temp_header_array = []
                        final_array = []

                        temp_header_array.append(WordElement(poster_code, self.code_header_style))
                        temp_header_array.append(WordElement("", self.code_header_style))
                        temp_header_array.append(WordElement("", self.code_header_style))

                        if self.HAS_RELEASE_IN_COVER:
                            temp_header_array.append(WordElement(film_info.release, self.release_date_style))
                        else:
                            temp_header_array.append(WordElement("", self.release_date_style))

                        final_array.append(temp_header_array)

                        Word.create_simple_table(document=self.document, parent=j,
                                                 table=table,
                                                 content=final_array,
                                                 empty_style=self.empty_style,
                                                 mode=4
                                                 )

                    if self.film_type == 0 or self.film_type == 2 or self.film_type == 3:
                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(film_info.eng_title, self.eng_title_header_style,
                                                                       alignment=1))

                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(film_info.per_title, self.per_title_header_style,
                                                                       alignment=1))

                    elif self.film_type == 1:
                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(film_info.per_title, self.per_title_header_style,
                                                                       alignment=1))

                    if self.film_type == 0 or self.film_type == 1 or self.film_type == 3:
                        per_genre = " , ".join(film_info.per_genre)
                        genre_text = "(" + per_genre + ")"
                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(genre_text, self.per_title_header_style,
                                                                       alignment=1))

                counter += 1

        self.document.save(self.final_word_location)

        if self.EXPORT_DOCX:
            os.system(self.final_word_location)

        if self.EXPORT_PDF:
            Word.convert_to_pdf(self.final_word_location, self.final_pdf_location)

    def poster_nine_part(self, films_array):
        self.set_margins()

        # if self.film_type == 1 or self.film_type == 2:
        #     self.MODE_8_PIC_HEIGHT = 4.80

        films_length = len(films_array)

        table1 = self.document.add_table(rows=math.ceil(films_length / 3), cols=3)

        counter = 0

        for i in table1.rows:
            for j in i.cells:
                if counter < films_length:
                    film = films_array[counter]

                    film_path = film[0]

                    Print.print_green(get_last_file_or_folder_from_path(film[0]))

                    pic_location = rename_file_paths_for_files(film[0] + "\\" + "Folder.jpg")
                    film_info = film[1]
                    poster_code = "(" + File.trim_poster_code(get_last_file_or_folder_from_path(film[0])) + ")"

                    # j.paragraphs[0].style = self.empty_style
                    Word.delete_paragraph(j.paragraphs[0])

                    if self.film_type == 0:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_9_PIC_WIDTH,
                                                    pic_height=self.MODE_9_PIC_HEIGHT_TYPE_0,
                                                    alignment=1)

                    elif self.film_type == 1:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_9_PIC_WIDTH,
                                                    pic_height=self.MODE_9_PIC_HEIGHT_TYPE_1,
                                                    alignment=1)

                    elif self.film_type == 2:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_9_PIC_WIDTH,
                                                    pic_height=self.MODE_9_PIC_HEIGHT_TYPE_2,
                                                    alignment=1)

                    elif self.film_type == 3:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_9_PIC_WIDTH,
                                                    pic_height=self.MODE_9_PIC_HEIGHT_TYPE_3,
                                                    alignment=1)

                    # ----------------------- Poster Code & Release Date in Table ----------------------- #
                    if self.film_type == 0 or self.film_type == 3:
                        table = TableElement(self.table_style, 1)

                        temp_header_array = []
                        final_array = []

                        is_subbed, is_dubbed = has_per_subtitle(film_path)

                        temp_header_array.append(WordElement(poster_code, self.code_header_style))
                        # temp_header_array.append(WordElement("", self.code_header_style))

                        score = film_info.score + "/10"

                        temp_header_array.append(
                            WordElement(score, self.score_style))

                        temp_header_array.append(
                            WordElement("زیرنویس", self.subbed_style, is_checkbox=1,
                                        check_state=is_subbed, width=self.MODE_9_dub_sub_cell_width))
                        temp_header_array.append(
                            WordElement("دوبله", self.dubbed_style, is_checkbox=1,
                                        check_state=is_dubbed, width=self.MODE_9_dub_sub_cell_width))

                        # temp_header_array.append(WordElement("", self.code_header_style))

                        if self.HAS_RELEASE_IN_COVER:
                            temp_header_array.append(WordElement(film_info.release, self.release_date_style))
                        else:
                            temp_header_array.append(WordElement("", self.release_date_style))

                        final_array.append(temp_header_array)

                        Word.create_simple_table(document=self.document, parent=j,
                                                 table=table,
                                                 content=final_array,
                                                 empty_style=self.empty_style,
                                                 mode=4
                                                 )

                        if self.film_type == 3:
                            Word.add_input_checkbox_content_control(parent=j,
                                                                    element=WordElement(
                                                                        "تمام شده", self.finished_style, is_checkbox=1,
                                                                        check_state=self.HAS_TV_SERIES_FINISHED_IN_COVER
                                                                    ))

                    if self.film_type == 1 or self.film_type == 2:
                        table = TableElement(self.table_style, 1)

                        temp_header_array = []
                        final_array = []

                        temp_header_array.append(WordElement(poster_code, self.code_header_style))
                        temp_header_array.append(WordElement("", self.code_header_style))
                        temp_header_array.append(WordElement("", self.code_header_style))

                        if self.HAS_RELEASE_IN_COVER:
                            temp_header_array.append(WordElement(film_info.release, self.release_date_style))
                        else:
                            temp_header_array.append(WordElement("", self.release_date_style))

                        final_array.append(temp_header_array)

                        Word.create_simple_table(document=self.document, parent=j,
                                                 table=table,
                                                 content=final_array,
                                                 empty_style=self.empty_style,
                                                 mode=4
                                                 )

                    if self.film_type == 0 or self.film_type == 2 or self.film_type == 3:
                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(film_info.eng_title, self.eng_title_header_style,
                                                                       alignment=1))

                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(film_info.per_title, self.per_title_header_style,
                                                                       alignment=1))


                    elif self.film_type == 1:
                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(film_info.per_title, self.per_title_header_style,
                                                                       alignment=1))

                    if self.film_type == 0 or self.film_type == 1 or self.film_type == 3:
                        per_genre = " , ".join(film_info.per_genre)
                        genre_text = "(" + per_genre + ")"
                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(genre_text, self.per_title_header_style,
                                                                       alignment=1))

                counter += 1

        self.document.save(self.final_word_location)

        if self.EXPORT_DOCX:
            os.system(self.final_word_location)

        if self.EXPORT_PDF:
            Word.convert_to_pdf(self.final_word_location, self.final_pdf_location)

    def poster_twelve_part(self, films_array):
        self.set_margins()

        # if self.film_type == 1 or self.film_type == 2:
        #     self.MODE_12_PIC_HEIGHT = 4.80

        films_length = len(films_array)

        table1 = self.document.add_table(rows=math.ceil(films_length / 4), cols=4)

        counter = 0

        for i in table1.rows:
            for j in i.cells:
                if counter < films_length:
                    film = films_array[counter]

                    film_path = film[0]

                    Print.print_green(get_last_file_or_folder_from_path(film[0]))

                    pic_location = rename_file_paths_for_files(film[0] + "\\" + "Folder.jpg")
                    film_info = film[1]
                    poster_code = "(" + File.trim_poster_code(get_last_file_or_folder_from_path(film[0])) + ")"

                    # j.paragraphs[0].style = self.empty_style
                    Word.delete_paragraph(j.paragraphs[0])

                    if self.film_type == 0:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_12_PIC_WIDTH,
                                                    pic_height=self.MODE_12_PIC_HEIGHT_TYPE_0,
                                                    alignment=1)

                    elif self.film_type == 1:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_12_PIC_WIDTH,
                                                    pic_height=self.MODE_12_PIC_HEIGHT_TYPE_1,
                                                    alignment=1)

                    elif self.film_type == 2:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_12_PIC_WIDTH,
                                                    pic_height=self.MODE_12_PIC_HEIGHT_TYPE_2,
                                                    alignment=1)

                    elif self.film_type == 3:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_12_PIC_WIDTH,
                                                    pic_height=self.MODE_12_PIC_HEIGHT_TYPE_3,
                                                    alignment=1)

                    # ----------------------- Poster Code & Release Date in Table ----------------------- #
                    if self.film_type == 0 or self.film_type == 3:
                        table = TableElement(self.table_style, 1)

                        temp_header_array = []
                        final_array = []

                        is_subbed, is_dubbed = has_per_subtitle(film_path)

                        temp_header_array.append(WordElement(poster_code, self.code_header_style))
                        # temp_header_array.append(WordElement("", self.code_header_style))

                        score = film_info.score + "/10"

                        temp_header_array.append(
                            WordElement(score, self.score_style))

                        temp_header_array.append(
                            WordElement("زیرنویس", self.subbed_style, is_checkbox=1,
                                        check_state=is_subbed, width=self.MODE_16_dub_sub_cell_width))
                        temp_header_array.append(
                            WordElement("دوبله", self.dubbed_style, is_checkbox=1,
                                        check_state=is_dubbed, width=self.MODE_16_dub_sub_cell_width))

                        # temp_header_array.append(WordElement("", self.code_header_style))

                        if self.HAS_RELEASE_IN_COVER:
                            temp_header_array.append(WordElement(film_info.release, self.release_date_style, width=25))
                        else:
                            temp_header_array.append(WordElement("", self.code_header_style, width=50))

                        final_array.append(temp_header_array)

                        Word.create_simple_table(document=self.document, parent=j,
                                                 table=table,
                                                 content=final_array,
                                                 empty_style=self.empty_style,
                                                 mode=4
                                                 )

                        if self.film_type == 3:
                            Word.add_input_checkbox_content_control(parent=j,
                                                                    element=WordElement(
                                                                        "تمام شده", self.finished_style, is_checkbox=1,
                                                                        check_state=self.HAS_TV_SERIES_FINISHED_IN_COVER
                                                                    ))

                    if self.film_type == 1 or self.film_type == 2:
                        table = TableElement(self.table_style, 1)

                        temp_header_array = []
                        final_array = []

                        temp_header_array.append(WordElement(poster_code, self.code_header_style))
                        temp_header_array.append(WordElement("", self.code_header_style))
                        temp_header_array.append(WordElement("", self.code_header_style))

                        if self.HAS_RELEASE_IN_COVER:
                            temp_header_array.append(WordElement(film_info.release, self.release_date_style))
                        else:
                            temp_header_array.append(WordElement("", self.release_date_style))

                        final_array.append(temp_header_array)

                        Word.create_simple_table(document=self.document, parent=j,
                                                 table=table,
                                                 content=final_array,
                                                 empty_style=self.empty_style,
                                                 mode=4
                                                 )

                    if self.film_type == 0 or self.film_type == 2 or self.film_type == 3:
                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(film_info.eng_title, self.eng_title_header_style,
                                                                       alignment=1))

                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(film_info.per_title, self.per_title_header_style,
                                                                       alignment=1))

                    elif self.film_type == 1:
                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(film_info.per_title, self.per_title_header_style,
                                                                       alignment=1))

                    if self.film_type == 0 or self.film_type == 1 or self.film_type == 3:
                        per_genre = " , ".join(film_info.per_genre)
                        genre_text = "(" + per_genre + ")"
                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(genre_text, self.per_title_header_style,
                                                                       alignment=1))

                counter += 1

        # Word.create_paragraph_text(parent=self.document, element=WordElement(content="", style=self.empty_style))

        self.document.save(self.final_word_location)

        if self.EXPORT_DOCX:
            os.system(self.final_word_location)

        if self.EXPORT_PDF:
            Word.convert_to_pdf(self.final_word_location, self.final_pdf_location)

    def poster_fifteen_part(self, films_array):
        self.set_margins()

        # if self.film_type == 1 or self.film_type == 2:
        #     self.MODE_14_PIC_HEIGHT = 4.80

        films_length = len(films_array)

        table1 = self.document.add_table(rows=math.ceil(films_length / 3), cols=5)

        counter = 0

        for i in table1.rows:
            for j in i.cells:
                if counter < films_length:
                    film = films_array[counter]

                    film_path = film[0]

                    Print.print_green(get_last_file_or_folder_from_path(film[0]))

                    pic_location = rename_file_paths_for_files(film[0] + "\\" + "Folder.jpg")
                    film_info = film[1]
                    poster_code = "(" + File.trim_poster_code(get_last_file_or_folder_from_path(film[0])) + ")"

                    # j.paragraphs[0].style = self.empty_style
                    Word.delete_paragraph(j.paragraphs[0])

                    if self.film_type == 0:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_15_PIC_WIDTH,
                                                    pic_height=self.MODE_15_PIC_HEIGHT_TYPE_0,
                                                    alignment=1)

                    elif self.film_type == 1:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_15_PIC_WIDTH,
                                                    pic_height=self.MODE_15_PIC_HEIGHT_TYPE_1,
                                                    alignment=1)

                    elif self.film_type == 2:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_15_PIC_WIDTH,
                                                    pic_height=self.MODE_15_PIC_HEIGHT_TYPE_2,
                                                    alignment=1)

                    elif self.film_type == 3:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_15_PIC_WIDTH,
                                                    pic_height=self.MODE_15_PIC_HEIGHT_TYPE_3,
                                                    alignment=1)

                    # ----------------------- Poster Code & Release Date in Table ----------------------- #
                    if self.film_type == 0 or self.film_type == 3:
                        table = TableElement(self.table_style, 1)

                        temp_header_array = []
                        final_array = []

                        is_subbed, is_dubbed = has_per_subtitle(film_path)

                        temp_header_array.append(WordElement(poster_code, self.code_header_style))
                        # temp_header_array.append(WordElement("", self.code_header_style))

                        score = film_info.score + "/10"

                        temp_header_array.append(
                            WordElement(score, self.score_style))

                        temp_header_array.append(
                            WordElement("زیرنویس", self.subbed_style, is_checkbox=1,
                                        check_state=is_subbed, width=self.MODE_16_dub_sub_cell_width))
                        temp_header_array.append(
                            WordElement("دوبله", self.dubbed_style, is_checkbox=1,
                                        check_state=is_dubbed, width=self.MODE_16_dub_sub_cell_width))

                        # temp_header_array.append(WordElement("", self.code_header_style))

                        if self.HAS_RELEASE_IN_COVER:
                            temp_header_array.append(WordElement(film_info.release, self.code_header_style, width=25))
                        else:
                            temp_header_array.append(WordElement("", self.code_header_style, width=50))

                        final_array.append(temp_header_array)

                        Word.create_simple_table(document=self.document, parent=j,
                                                 table=table,
                                                 content=final_array,
                                                 empty_style=self.empty_style,
                                                 mode=4
                                                 )

                        if self.film_type == 3:
                            Word.add_input_checkbox_content_control(parent=j,
                                                                    element=WordElement(
                                                                        "تمام شده", self.finished_style, is_checkbox=1,
                                                                        check_state=self.HAS_TV_SERIES_FINISHED_IN_COVER
                                                                    ))

                    if self.film_type == 1 or self.film_type == 2:
                        table = TableElement(self.table_style, 1)

                        temp_header_array = []
                        final_array = []

                        temp_header_array.append(WordElement(poster_code, self.code_header_style))
                        temp_header_array.append(WordElement("", self.code_header_style))
                        temp_header_array.append(WordElement("", self.code_header_style))

                        if self.HAS_RELEASE_IN_COVER:
                            temp_header_array.append(WordElement(film_info.release, self.release_date_style))
                        else:
                            temp_header_array.append(WordElement("", self.release_date_style))

                        final_array.append(temp_header_array)

                        Word.create_simple_table(document=self.document, parent=j,
                                                 table=table,
                                                 content=final_array,
                                                 empty_style=self.empty_style,
                                                 mode=4
                                                 )

                    if self.film_type == 0 or self.film_type == 2 or self.film_type == 3:
                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(film_info.eng_title, self.eng_title_header_style,
                                                                       alignment=1))

                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(film_info.per_title, self.per_title_header_style,
                                                                       alignment=1))

                    elif self.film_type == 1:
                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(film_info.per_title, self.per_title_header_style,
                                                                       alignment=1))

                    if self.film_type == 0 or self.film_type == 1 or self.film_type == 3:
                        per_genre = " , ".join(film_info.per_genre)
                        genre_text = "(" + per_genre + ")"
                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(genre_text, self.per_title_header_style,
                                                                       alignment=1))

                counter += 1

        # Word.create_paragraph_text(parent=self.document, element=WordElement(content="", style=self.empty_style))

        self.document.save(self.final_word_location)

        if self.EXPORT_DOCX:
            os.system(self.final_word_location)

        if self.EXPORT_PDF:
            Word.convert_to_pdf(self.final_word_location, self.final_pdf_location)

    def poster_sixteen_part(self, films_array):
        self.set_margins()

        # if self.film_type == 1 or self.film_type == 2:
        #     self.MODE_16_PIC_HEIGHT = 4.80

        films_length = len(films_array)

        table1 = self.document.add_table(rows=math.ceil(films_length / 4), cols=4)

        counter = 0

        for i in table1.rows:
            for j in i.cells:
                if counter < films_length:
                    film = films_array[counter]

                    film_path = film[0]

                    Print.print_green(get_last_file_or_folder_from_path(film[0]))

                    pic_location = rename_file_paths_for_files(film[0] + "\\" + "Folder.jpg")
                    film_info = film[1]
                    poster_code = "(" + File.trim_poster_code(get_last_file_or_folder_from_path(film[0])) + ")"

                    # j.paragraphs[0].style = self.empty_style
                    Word.delete_paragraph(j.paragraphs[0])

                    if self.film_type == 0:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_16_PIC_WIDTH,
                                                    pic_height=self.MODE_16_PIC_HEIGHT_TYPE_0,
                                                    alignment=1)

                    elif self.film_type == 1:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_16_PIC_WIDTH,
                                                    pic_height=self.MODE_16_PIC_HEIGHT_TYPE_1,
                                                    alignment=1)

                    elif self.film_type == 2:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_16_PIC_WIDTH,
                                                    pic_height=self.MODE_16_PIC_HEIGHT_TYPE_2,
                                                    alignment=1)

                    elif self.film_type == 3:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_16_PIC_WIDTH,
                                                    pic_height=self.MODE_16_PIC_HEIGHT_TYPE_3,
                                                    alignment=1)

                    # ----------------------- Poster Code & Release Date in Table ----------------------- #
                    if self.film_type == 0 or self.film_type == 3:
                        table = TableElement(self.table_style, 1)

                        temp_header_array = []
                        final_array = []

                        is_subbed, is_dubbed = has_per_subtitle(film_path)

                        temp_header_array.append(WordElement(poster_code, self.code_header_style))
                        # temp_header_array.append(WordElement("", self.code_header_style))

                        score = film_info.score + "/10"

                        temp_header_array.append(
                            WordElement(score, self.score_style))

                        temp_header_array.append(
                            WordElement("زیرنویس", self.subbed_style, is_checkbox=1,
                                        check_state=is_subbed, width=self.MODE_16_dub_sub_cell_width))
                        temp_header_array.append(
                            WordElement("دوبله", self.dubbed_style, is_checkbox=1,
                                        check_state=is_dubbed, width=self.MODE_16_dub_sub_cell_width))

                        # temp_header_array.append(WordElement("", self.code_header_style))

                        if self.HAS_RELEASE_IN_COVER:
                            temp_header_array.append(WordElement(film_info.release, self.code_header_style, width=25))
                        else:
                            temp_header_array.append(WordElement("", self.code_header_style, width=50))

                        final_array.append(temp_header_array)

                        Word.create_simple_table(document=self.document, parent=j,
                                                 table=table,
                                                 content=final_array,
                                                 empty_style=self.empty_style,
                                                 mode=4
                                                 )

                        if self.film_type == 3:
                            Word.add_input_checkbox_content_control(parent=j,
                                                                    element=WordElement(
                                                                        "تمام شده", self.finished_style, is_checkbox=1,
                                                                        check_state=self.HAS_TV_SERIES_FINISHED_IN_COVER
                                                                    ))

                    if self.film_type == 1 or self.film_type == 2:
                        table = TableElement(self.table_style, 1)

                        temp_header_array = []
                        final_array = []

                        temp_header_array.append(WordElement(poster_code, self.code_header_style))
                        temp_header_array.append(WordElement("", self.code_header_style))
                        temp_header_array.append(WordElement("", self.code_header_style))

                        if self.HAS_RELEASE_IN_COVER:
                            temp_header_array.append(WordElement(film_info.release, self.release_date_style))
                        else:
                            temp_header_array.append(WordElement("", self.release_date_style))

                        final_array.append(temp_header_array)

                        Word.create_simple_table(document=self.document, parent=j,
                                                 table=table,
                                                 content=final_array,
                                                 empty_style=self.empty_style,
                                                 mode=4
                                                 )

                    if self.film_type == 0 or self.film_type == 2 or self.film_type == 3:
                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(film_info.eng_title, self.eng_title_header_style,
                                                                       alignment=1))

                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(film_info.per_title, self.per_title_header_style,
                                                                       alignment=1))

                    elif self.film_type == 1:
                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(film_info.per_title, self.per_title_header_style,
                                                                       alignment=1))

                    if self.film_type == 0 or self.film_type == 1 or self.film_type == 3:
                        per_genre = " , ".join(film_info.per_genre)
                        genre_text = "(" + per_genre + ")"
                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(genre_text, self.per_title_header_style,
                                                                       alignment=1))

                counter += 1

        # Word.create_paragraph_text(parent=self.document, element=WordElement(content="", style=self.empty_style))

        self.document.save(self.final_word_location)

        if self.EXPORT_DOCX:
            os.system(self.final_word_location)

        if self.EXPORT_PDF:
            Word.convert_to_pdf(self.final_word_location, self.final_pdf_location)

    def poster_eighteen_part(self, films_array):
        self.set_margins()

        # if self.film_type == 1 or self.film_type == 2:
        #     self.MODE_18_PIC_HEIGHT = 4.80

        films_length = len(films_array)

        table1 = self.document.add_table(rows=math.ceil(films_length / 3), cols=6)

        counter = 0

        for i in table1.rows:
            for j in i.cells:
                if counter < films_length:
                    film = films_array[counter]

                    film_path = film[0]

                    Print.print_green(get_last_file_or_folder_from_path(film[0]))

                    pic_location = rename_file_paths_for_files(film[0] + "\\" + "Folder.jpg")
                    film_info = film[1]
                    poster_code = "(" + File.trim_poster_code(get_last_file_or_folder_from_path(film[0])) + ")"

                    # j.paragraphs[0].style = self.empty_style
                    Word.delete_paragraph(j.paragraphs[0])

                    if self.film_type == 0:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_18_PIC_WIDTH,
                                                    pic_height=self.MODE_18_PIC_HEIGHT_TYPE_0,
                                                    alignment=1)

                    elif self.film_type == 1:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_18_PIC_WIDTH,
                                                    pic_height=self.MODE_18_PIC_HEIGHT_TYPE_1,
                                                    alignment=1)

                    elif self.film_type == 2:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_18_PIC_WIDTH,
                                                    pic_height=self.MODE_18_PIC_HEIGHT_TYPE_2,
                                                    alignment=1)

                    elif self.film_type == 3:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_18_PIC_WIDTH,
                                                    pic_height=self.MODE_18_PIC_HEIGHT_TYPE_3,
                                                    alignment=1)

                    # ----------------------- Poster Code & Release Date in Table ----------------------- #
                    if self.film_type == 0 or self.film_type == 3:
                        table = TableElement(self.table_style, 1)

                        temp_header_array = []
                        final_array = []

                        is_subbed, is_dubbed = has_per_subtitle(film_path)

                        temp_header_array.append(WordElement(poster_code, self.code_header_style))
                        # temp_header_array.append(WordElement("", self.code_header_style))

                        score = film_info.score + "/10"

                        temp_header_array.append(
                            WordElement(score, self.score_style))

                        temp_header_array.append(
                            WordElement("زیرنویس", self.subbed_style, is_checkbox=1,
                                        check_state=is_subbed, width=self.MODE_16_dub_sub_cell_width))
                        temp_header_array.append(
                            WordElement("دوبله", self.dubbed_style, is_checkbox=1,
                                        check_state=is_dubbed, width=self.MODE_16_dub_sub_cell_width))

                        # temp_header_array.append(WordElement("", self.code_header_style))

                        if self.HAS_RELEASE_IN_COVER:
                            temp_header_array.append(WordElement(film_info.release, self.code_header_style, width=25))
                        else:
                            temp_header_array.append(WordElement("", self.code_header_style, width=50))

                        final_array.append(temp_header_array)

                        Word.create_simple_table(document=self.document, parent=j,
                                                 table=table,
                                                 content=final_array,
                                                 empty_style=self.empty_style,
                                                 mode=4
                                                 )

                        if self.film_type == 3:
                            Word.add_input_checkbox_content_control(parent=j,
                                                                    element=WordElement(
                                                                        "تمام شده", self.finished_style, is_checkbox=1,
                                                                        check_state=self.HAS_TV_SERIES_FINISHED_IN_COVER
                                                                    ))

                    if self.film_type == 1 or self.film_type == 2:
                        table = TableElement(self.table_style, 1)

                        temp_header_array = []
                        final_array = []

                        temp_header_array.append(WordElement(poster_code, self.code_header_style))
                        temp_header_array.append(WordElement("", self.code_header_style))
                        temp_header_array.append(WordElement("", self.code_header_style))

                        if self.HAS_RELEASE_IN_COVER:
                            temp_header_array.append(WordElement(film_info.release, self.release_date_style))
                        else:
                            temp_header_array.append(WordElement("", self.release_date_style))

                        final_array.append(temp_header_array)

                        Word.create_simple_table(document=self.document, parent=j,
                                                 table=table,
                                                 content=final_array,
                                                 empty_style=self.empty_style,
                                                 mode=4
                                                 )

                    if self.film_type == 0 or self.film_type == 2 or self.film_type == 3:
                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(film_info.eng_title, self.eng_title_header_style,
                                                                       alignment=1))

                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(film_info.per_title, self.per_title_header_style,
                                                                       alignment=1))

                    elif self.film_type == 1:
                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(film_info.per_title, self.per_title_header_style,
                                                                       alignment=1))

                    if self.film_type == 0 or self.film_type == 1 or self.film_type == 3:
                        per_genre = " , ".join(film_info.per_genre)
                        genre_text = "(" + per_genre + ")"
                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(genre_text, self.per_title_header_style,
                                                                       alignment=1))

                counter += 1

        # Word.create_paragraph_text(parent=self.document, element=WordElement(content="", style=self.empty_style))

        self.document.save(self.final_word_location)

        if self.EXPORT_DOCX:
            os.system(self.final_word_location)

        if self.EXPORT_PDF:
            Word.convert_to_pdf(self.final_word_location, self.final_pdf_location)

    def poster_twenty_part(self, films_array):
        self.set_margins()

        # if self.film_type == 1 or self.film_type == 2:
        #     self.MODE_20_PIC_HEIGHT = 4.80

        films_length = len(films_array)

        table1 = self.document.add_table(rows=math.ceil(films_length / 5), cols=5)

        counter = 0

        for i in table1.rows:
            for j in i.cells:
                if counter < films_length:
                    film = films_array[counter]

                    film_path = film[0]

                    Print.print_green(get_last_file_or_folder_from_path(film[0]))

                    pic_location = rename_file_paths_for_files(film[0] + "\\" + "Folder.jpg")
                    film_info = film[1]
                    poster_code = "(" + File.trim_poster_code(get_last_file_or_folder_from_path(film[0])) + ")"

                    # j.paragraphs[0].style = self.empty_style
                    Word.delete_paragraph(j.paragraphs[0])

                    if self.film_type == 0:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_20_PIC_WIDTH,
                                                    pic_height=self.MODE_20_PIC_HEIGHT_TYPE_0,
                                                    alignment=1)

                    elif self.film_type == 1:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_20_PIC_WIDTH,
                                                    pic_height=self.MODE_20_PIC_HEIGHT_TYPE_1,
                                                    alignment=1)

                    elif self.film_type == 2:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_20_PIC_WIDTH,
                                                    pic_height=self.MODE_20_PIC_HEIGHT_TYPE_2,
                                                    alignment=1)

                    elif self.film_type == 3:
                        Word.create_paragraph_image(parent=j, style=self.eng_title_header_style, pic_dir=pic_location,
                                                    pic_width=self.MODE_20_PIC_WIDTH,
                                                    pic_height=self.MODE_20_PIC_HEIGHT_TYPE_3,
                                                    alignment=1)

                    # ----------------------- Poster Code & Release Date in Table ----------------------- #
                    if self.film_type == 0 or self.film_type == 3:
                        table = TableElement(self.table_style, 1)

                        temp_header_array = []
                        final_array = []

                        is_subbed, is_dubbed = has_per_subtitle(film_path)

                        temp_header_array.append(WordElement(poster_code, self.code_header_style))
                        # temp_header_array.append(WordElement("", self.code_header_style))

                        score = film_info.score + "/10"

                        temp_header_array.append(
                            WordElement(score, self.score_style))

                        temp_header_array.append(
                            WordElement("زیرنویس", self.subbed_style, is_checkbox=1,
                                        check_state=is_subbed, width=self.MODE_20_dub_sub_cell_width))
                        temp_header_array.append(
                            WordElement("دوبله", self.dubbed_style, is_checkbox=1,
                                        check_state=is_dubbed, width=self.MODE_20_dub_sub_cell_width))

                        # temp_header_array.append(WordElement("", self.code_header_style))

                        if self.HAS_RELEASE_IN_COVER:
                            temp_header_array.append(WordElement(film_info.release, self.release_date_style))
                        else:
                            temp_header_array.append(WordElement("", self.release_date_style))

                        final_array.append(temp_header_array)

                        Word.create_simple_table(document=self.document, parent=j,
                                                 table=table,
                                                 content=final_array,
                                                 empty_style=self.empty_style,
                                                 mode=4
                                                 )

                        if self.film_type == 3:
                            Word.add_input_checkbox_content_control(parent=j,
                                                                    element=WordElement(
                                                                        "تمام شده", self.finished_style, is_checkbox=1,
                                                                        check_state=self.HAS_TV_SERIES_FINISHED_IN_COVER
                                                                    ))

                    if self.film_type == 1 or self.film_type == 2:
                        table = TableElement(self.table_style, 1)

                        temp_header_array = []
                        final_array = []

                        temp_header_array.append(WordElement(poster_code, self.code_header_style))
                        temp_header_array.append(WordElement("", self.code_header_style))
                        temp_header_array.append(WordElement("", self.code_header_style))

                        if self.HAS_RELEASE_IN_COVER:
                            temp_header_array.append(WordElement(film_info.release, self.release_date_style))
                        else:
                            temp_header_array.append(WordElement("", self.release_date_style))

                        final_array.append(temp_header_array)

                        Word.create_simple_table(document=self.document, parent=j,
                                                 table=table,
                                                 content=final_array,
                                                 empty_style=self.empty_style,
                                                 mode=4
                                                 )

                    if self.film_type == 0 or self.film_type == 2 or self.film_type == 3:
                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(film_info.eng_title, self.eng_title_header_style,
                                                                       alignment=1))

                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(film_info.per_title, self.per_title_header_style,
                                                                       alignment=1))

                    elif self.film_type == 1:
                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(film_info.per_title, self.per_title_header_style,
                                                                       alignment=1))

                    if self.film_type == 0 or self.film_type == 1 or self.film_type == 3:
                        per_genre = " , ".join(film_info.per_genre)
                        genre_text = "(" + per_genre + ")"
                        Word.create_paragraph_text(parent=j,
                                                   element=WordElement(genre_text, self.per_title_header_style,
                                                                       alignment=1))

                counter += 1

        # Word.create_paragraph_text(parent=self.document, element=WordElement(content="", style=self.empty_style))

        self.document.save(self.final_word_location)
        os.system(self.final_word_location)
        # Word.convert_to_pdf(self.final_word_location, self.full_path, self.POSTER_PDF_NAME)
