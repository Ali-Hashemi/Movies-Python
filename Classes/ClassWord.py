from docx import Document, oxml
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor
from docx2pdf import convert
import numpy as np
from Classes.ClassUtility import *
from docx2pdf import convert
from Data.WordElement import *
from Data.TableElement import *
import aspose.words as aw
import comtypes.client
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn


class Word:
    @staticmethod
    def create_paragraph_text(parent, element: WordElement):
        p = parent.add_paragraph(style=element.style)
        Word.get_paragraph_format(p)
        p.alignment = element.alignment
        r = p.add_run()
        r.add_text(element.content)

    @staticmethod
    def get_paragraph_format(paragraph, space_before=0, space_after=0):
        format = paragraph.paragraph_format
        format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        format.space_before = Pt(space_before)
        format.space_after = Pt(space_after)
        return format

    @staticmethod
    def create_paragraph_image(parent, pic_dir, pic_width, pic_height, style=None, alignment=1):
        # p.alignment = 3  # for left, 1 for center, 2 right, 3 justify ....
        if style:
            p = parent.add_paragraph(style=style)
        else:
            p = parent.add_paragraph()
        Word.get_paragraph_format(p, space_after=5)
        p.alignment = alignment
        r = p.add_run()
        r.add_picture(pic_dir, width=Inches(pic_width), height=Inches(pic_height))

    @staticmethod
    def create_simple_table(document, parent,
                            table,
                            content,
                            empty_style=None,
                            mode=None,
                            film_type=None
                            ):
        # -------------- Example of content ----------------#
        # columns_values = [
        #     [
        #         WordElement(poster_code, self.code_header_style),
        #         WordElement(poster_code, self.code_header_style),
        #         WordElement(poster_code, self.code_header_style)
        #     ]
        # ]
        # --------------------------------------------------#

        row_number = len(content)

        for index, i in enumerate(content):
            column_number = len(content[index])
            break

        created_table = parent.add_table(rows=row_number, cols=column_number)

        if mode == 2:
            Word.delete_paragraph(parent.paragraphs[0])
            Word.delete_paragraph(parent.paragraphs[0])
        elif mode == 4:
            Word.delete_paragraph(parent.paragraphs[1])

        created_table.alignment = table.alignment
        created_table.style = table.style

        # for index_row, item in enumerate(content):

        for index_cell in range(column_number):
            current_row = created_table.rows[0]

            current_cell = current_row.cells[index_cell]

            Word.set_cell_margins(current_cell, top=0, start=30, bottom=0, end=30)

            current_element = content[0][index_cell]

            if not current_element.is_checkbox:
                current_cell.text = current_element.content
                current_cell.paragraphs[0].style = current_element.style
                current_cell.paragraphs[0].alignment = current_element.alignment
            else:
                # current_cell.paragraphs[0].style = empty_style
                Word.delete_paragraph(current_cell.paragraphs[0])

                Word.add_input_checkbox_content_control(parent=current_cell,
                                                        element=current_element)
                if current_element.width:
                    current_cell.width = Inches(current_element.width)
                if current_element.height:
                    current_cell.height = Inches(current_element.height)

    @staticmethod
    def create_table(parent, rows, cols,
                     table_styles,
                     table_cel1_style,
                     table_cel2_style,
                     empty_style,
                     score_cell_styles,
                     content,
                     cell1_width, cell1_height,
                     cell2_width):

        joiner_character = ' ، '

        table = parent.add_table(rows=rows, cols=cols)
        table.style = table_styles

        for index, item in enumerate(content):
            row1 = table.rows[index]

            row1.height = Inches(cell1_height)

            cell1 = table.rows[index].cells[1]
            cell1.text = item
            cell1.paragraphs[0].style = table_cel1_style
            cell1.width = Inches(cell1_width)

            cell2 = table.rows[index].cells[0]

            if isinstance(content[item], (list, tuple, np.ndarray)):
                if find_in_text(item, "بازیگر"):
                    actors_list = []
                    counter = 0
                    for n in content[item]:
                        if counter < 4:
                            actors_list.append(n)
                            counter += 1
                        else:
                            break

                    final = joiner_character.join(actors_list)
                else:
                    final = joiner_character.join(content[item])

                if final:
                    cell2.text = final

                cell2.paragraphs[0].style = table_cel2_style

            elif find_in_text(item, "امتیاز"):
                score_values = [
                    "10",
                    "از",
                    content[item]
                ]

                Word.create_table_for_score(cell2, rows=1, cols=3,
                                            table_style=table_styles,
                                            width=0.30, values=score_values,
                                            cell_styles=score_cell_styles,
                                            alignment=2)

                cell2.paragraphs[0].style = empty_style

            else:
                content[item] = str(content[item]).replace('( ', '"')
                content[item] = str(content[item]).replace('(', '"')
                content[item] = str(content[item]).replace(' )', '"')
                content[item] = str(content[item]).replace(')', '"')

                text_content = str(content[item])

                if len(text_content) > 360:
                    cell2.text = text_content[:360]
                else:
                    cell2.text = text_content

                cell2.paragraphs[0].style = table_cel2_style

            cell2.width = Inches(cell2_width)

    @staticmethod
    def create_table_for_score(parent, rows, cols, table_style, width, values, cell_styles, alignment=1):
        table = parent.add_table(rows=rows, cols=cols)
        table.style = table_style
        table.alignment = alignment

        for row in table.rows:
            for index, cell in enumerate(row.cells):
                cell.width = Inches(width)
                cell.text = values[index]
                cell.paragraphs[0].style = cell_styles[index]

    @staticmethod
    def add_input_checkbox_content_control(parent, element: WordElement):
        paragraph = parent.add_paragraph(style=element.style)
        paragraph.alignment = element.alignment
        paragraph_tag = paragraph._element
        sdt = oxml.shared.OxmlElement('w:sdt')
        sdt_props = oxml.shared.OxmlElement('w:sdtPr')

        w_checkbox = oxml.shared.OxmlElement('w14:checkbox')
        w_checked = oxml.shared.OxmlElement('w14:checked')
        w_checked.set(oxml.ns.qn('w14:val'), "1" if element.check_state else "0")
        w_checked_state = oxml.shared.OxmlElement('w14:checkedState')
        w_checked_state.set(oxml.ns.qn('w14:font'), "MS Gothic")
        w_checked_state.set(oxml.ns.qn('w14:val'), "2611")  # unicode value for box with x in it

        w_unchecked_state = oxml.shared.OxmlElement('w14:uncheckedState')
        w_unchecked_state.set(oxml.ns.qn('w14:font'), "MS Gothic")
        w_unchecked_state.set(oxml.ns.qn('w14:val'), "2610")  # unicode value for empty box

        w_checkbox.append(w_checked)
        w_checkbox.append(w_checked_state)
        w_checkbox.append(w_unchecked_state)
        sdt_props.append(w_checkbox)

        sdt_content = oxml.shared.OxmlElement('w:sdtContent')
        sdt_content_run = oxml.OxmlElement('w:r')
        # the box doesn't appear until clicked, so a default unicode box will need to be set as a placeholder
        # sdt_content_run.text = "\u2612" if checked else "\u2610"

        sdt_content_run.text = "\u2611" + " " + str(element.content) if element.check_state else "\u2610" + " " + str(
            element.content)

        sdt_content.append(sdt_content_run)

        sdt.append(sdt_props)
        sdt.append(sdt_content)
        paragraph_tag.append(sdt)

    @staticmethod
    def create_empty_paragraph(parent, style):
        Word.create_paragraph_text(parent=parent, element=WordElement(content="", style=style))

    @staticmethod
    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None

    @staticmethod
    def convert_to_pdf(word_path, pdf_save_dir):
        # if os.path.exists(word_path):
        #     pdf_path = rename_file_paths_by_os(pdf_save_dir) + pdf_name + ".pdf"
        #     convert(word_path, pdf_path)
        #     os.system(pdf_save_dir + pdf_name + ".pdf")

        if os.path.exists(word_path):
            pdf_full_path = rename_file_paths_for_files(pdf_save_dir)
            # convert(word_path, pdf_path)
            # os.system(pdf_path)

            wdFormatPDF = 17
            in_file = os.path.abspath(word_path)
            out_file = os.path.abspath(pdf_full_path)

            word = comtypes.client.CreateObject('Word.Application')
            doc = word.Documents.Open(in_file)
            doc.SaveAs(out_file, FileFormat=wdFormatPDF)
            doc.Close()
            word.Quit()

            # convert_pdf_to_image(pdf_full_path, rename_file_paths_by_os("D:\Images"))

            os.system(pdf_full_path)

    @staticmethod
    def set_cell_margins(cell, **kwargs):
        """
        cell:  actual cell instance you want to modify
        usage:
            set_cell_margins(cell, top=50, start=50, bottom=50, end=50)

        provided values are in twentieths of a point (1/1440 of an inch).
        read more here: http://officeopenxml.com/WPtableCellMargins.php
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')

        for m in ["top", "start", "bottom", "end"]:
            if m in kwargs:
                node = OxmlElement("w:{}".format(m))
                node.set(qn('w:w'), str(kwargs.get(m)))
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)

        tcPr.append(tcMar)
